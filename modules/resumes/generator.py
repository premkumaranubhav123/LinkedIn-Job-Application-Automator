import docx
from fpdf import FPDF

def create_resume_docx(user_details, summary, experience, projects, skills, certificates):
    doc = docx.Document()
    doc.add_heading(user_details['name'], 0)
    doc.add_paragraph(user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'])
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)
    doc.add_heading('Experience', 1)
    for experience_item in experience:
        doc.add_heading(experience_item['company'], 2)
        doc.add_paragraph(experience_item['role'] + ' | ' + experience_item['dates'])
        doc.add_paragraph(experience_item['achievements'])
    doc.add_heading('Projects', 1)
    for project in projects:
        doc.add_heading(project['name'], 2)
        doc.add_paragraph(project['description'] + ' | ' + project['technologies'])
    doc.add_heading('Skills', 1)
    doc.add_paragraph(', '.join(skills))
    doc.add_heading('Certificates', 1)
    for certificate in certificates:
        doc.add_heading(certificate['name'], 2)
        doc.add_paragraph(certificate['description'])
    doc.save('resume.docx')
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)
    pdf.cell(0, 10, user_details['name'], 0, 1, 'C')
    pdf.cell(0, 10, user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'], 0, 1, 'L')
    pdf.cell(0, 10, 'Summary', 0, 1, 'L')
    pdf.multi_cell(0, 10, summary)
    pdf.cell(0, 10, 'Experience', 0, 1, 'L')
    for experience_item in experience:
        pdf.cell(0, 10, experience_item['company'], 0, 1, 'L')
        pdf.cell(0, 10, experience_item['role'] + ' | ' + experience_item['dates'], 0, 1, 'L')
        pdf.multi_cell(0, 10, experience_item['achievements'])
    pdf.cell(0, 10, 'Projects', 0, 1, 'L')
    for project in projects:
        pdf.cell(0, 10, project['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, project['description'] + ' | ' + project['technologies'])
    pdf.cell(0, 10, 'Skills', 0, 1, 'L')
    pdf.multi_cell(0, 10, ', '.join(skills))
    pdf.cell(0, 10, 'Certificates', 0, 1, 'L')
    for certificate in certificates:
        pdf.cell(0, 10, certificate['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, certificate['description'])
    pdf.output('resume.pdf', 'F')